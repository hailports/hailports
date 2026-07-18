"""Document output — saves files to OneDrive Claude Outputs folder.
Generates HTML from LLM, auto-converts to docx/pptx when requested."""

import asyncio
import os
import re
import logging
from datetime import datetime
from tools.base import BaseTool, make_tool_def
from core import SETTINGS, BRANDING
from core.doc_branding import get_doc_preferences

log = logging.getLogger(__name__)

_COLORS = BRANDING.get("colors", {})
_PRIMARY = _COLORS.get("primary", "#1B3A5C")
_PRIMARY_DARK = _COLORS.get("primary_dark", "#2d5280")
_ACCENT = _COLORS.get("accent", "#C5A55A")
_SF_BRAND_LABEL = BRANDING.get("tenant", {}).get("salesforce_display_name", "VALLEY")


def _sanitize_output_filename(filename: str, fmt: str = "html") -> str:
    """Normalize LLM-proposed filenames to something filesystem- and UI-safe."""
    filename = str(filename or "Document").strip()
    filename = filename.replace("\r", " ").replace("\n", " ")
    filename = filename.replace("—", "-").replace("–", "-")
    filename = filename.replace("“", '"').replace("”", '"')
    filename = filename.replace("‘", "'").replace("’", "'")
    filename = re.sub(r"[\x00-\x1f\x7f]", "", filename)
    filename = filename.replace("/", "-").replace("\\", "-").replace(":", " -")
    filename = re.sub(r'[<>"|?*]', '', filename)
    filename = re.sub(r"\s+", " ", filename).strip(" .-")
    if not filename:
        filename = "Document"
    root, ext = os.path.splitext(filename)
    ext = (ext or f".{fmt}").lower()
    if not root:
        root = "Document"
    return f"{root}{ext}"

_redacted_CSS = """
@page { size: letter; margin: 0.75in; }
* { box-sizing: border-box; }
body {
  font-family: 'Segoe UI', 'Helvetica Neue', Arial, sans-serif;
  color: #2c2c2c; line-height: 1.55; margin: 0; padding: 40px; background: #fff;
}
h1, h2, h3 { color: #1B3A5C; margin-top: 0; }
h1 { font-size: 34px; letter-spacing: -0.5px; }
h2 { font-size: 22px; border-bottom: 2px solid #C5A55A; padding-bottom: 6px; margin-top: 40px; }
h3 { font-size: 16px; margin-top: 24px; }
p { margin: 8px 0; }
ul, ol { margin: 8px 0 16px 24px; } li { margin-bottom: 6px; font-size: 14px; }

/* Cover */
.cover {
  background: linear-gradient(135deg, #1B3A5C 0%, #2d5280 100%);
  color: #fff; padding: 80px 60px; border-radius: 8px; margin-bottom: 40px; page-break-after: always;
}
.cover h1 { color: #fff; font-size: 44px; margin-bottom: 8px; }
.cover .subtitle { font-size: 20px; color: #C5A55A; margin-bottom: 60px; font-weight: 300; }
.cover .meta { font-size: 14px; color: #e0e0e0; line-height: 1.9; border-top: 1px solid rgba(255,255,255,0.2); padding-top: 20px; margin-top: 80px; }
.cover .meta strong { color: #C5A55A; }

/* Callout */
.callout {
  background: #F5F5F7; border-left: 4px solid #C5A55A; padding: 14px 18px;
  margin: 16px 0; border-radius: 4px; font-size: 14px;
}
.callout strong { color: #1B3A5C; }

/* Flow diagram */
.flow { display: flex; align-items: center; justify-content: space-between; gap: 10px; margin: 30px 0; flex-wrap: nowrap; }
.flow-step {
  flex: 1; background: #fff; border: 2px solid #1B3A5C; border-radius: 8px;
  padding: 16px 10px; text-align: center; font-size: 12px; font-weight: 600;
  color: #1B3A5C; box-shadow: 0 2px 4px rgba(0,0,0,0.06); position: relative;
}
.flow-step.start { background: #1B3A5C; color: #fff; border-radius: 40px; }
.flow-step.end { background: #C5A55A; color: #fff; border-radius: 40px; border-color: #C5A55A; }
.flow-step.decision { background: #fff8e8; border-color: #C5A55A; }
.flow-step .num { display: block; font-size: 10px; color: #888; margin-bottom: 4px; font-weight: 400; }
.flow-arrow { color: #1B3A5C; font-size: 24px; font-weight: bold; flex: 0 0 auto; }

/* Mockup — Salesforce Lightning UI */
.mockup {
  border: 1px solid #d0d0d5; border-radius: 6px; overflow: hidden; margin: 20px 0;
  box-shadow: 0 4px 12px rgba(0,0,0,0.08); font-family: 'Segoe UI', sans-serif; background: #fff;
}
.sf-browser-bar {
  background: #e8e8ed; padding: 8px 12px; display: flex; align-items: center; gap: 6px;
  border-bottom: 1px solid #d0d0d5;
}
.sf-browser-bar .dot { width: 10px; height: 10px; border-radius: 50%; background: #ccc; }
.sf-browser-bar .dot.red { background: #ff5f57; }
.sf-browser-bar .dot.yellow { background: #febc2e; }
.sf-browser-bar .dot.green { background: #28c840; }
.sf-browser-bar .url {
  background: #fff; flex: 1; margin-left: 10px; padding: 3px 10px; border-radius: 4px;
  font-size: 11px; color: #666; border: 1px solid #d0d0d5;
}
/* Valley logo strip */
.sf-valley-bar {
  background: #fff; display: flex; align-items: center; padding: 4px 16px; height: 30px;
  border-bottom: 1px solid #e0e0e0;
}
.sf-valley-logo { font-weight: 900; font-size: 15px; letter-spacing: 2px; font-family: Arial Black, Arial, sans-serif; color: #1a1a2e; }
.sf-valley-bar .sf-search { background: #f4f6f9; border: 1px solid #d8dde6; border-radius: 4px; padding: 3px 10px; color: #333; font-size: 11px; width: 220px; margin: 0 auto; }
.sf-valley-bar .sf-right { display: flex; align-items: center; gap: 8px; flex-shrink: 0; }
.sf-valley-bar .sf-search::placeholder { color: #999; }
.sf-valley-bar .sf-utils { display: flex; gap: 8px; align-items: center; font-size: 14px; color: #54698d; }
.sf-valley-bar .sf-avatar { width: 24px; height: 24px; border-radius: 50%; background: #032e61; display: flex; align-items: center; justify-content: center; font-size: 10px; font-weight: 700; color: #fff; position: relative; }
.sf-valley-bar .sf-avatar .notif { position: absolute; top: -4px; right: -4px; min-width: 14px; height: 14px; background: #e8112d; border-radius: 7px; font-size: 8px; display: flex; align-items: center; justify-content: center; border: 2px solid #fff; color: #fff; padding: 0 2px; }
/* SF Navigation bar */
.sf-header {
  background: #032e61; color: #fff; padding: 0 12px; display: flex;
  align-items: center; font-size: 12px; height: 38px;
  font-family: 'Salesforce Sans', 'Segoe UI', system-ui, sans-serif;
}
.sf-header .waffle { display: inline-grid; grid-template-columns: repeat(3,4px); gap: 2px; margin-right: 10px; cursor: pointer; }
.sf-header .waffle span { width: 4px; height: 4px; background: #fff; border-radius: 1px; }
.sf-header .app-name { font-weight: 600; font-size: 13px; margin-right: 16px; white-space: nowrap; }
.sf-header .nav { display: flex; gap: 0; font-size: 11px; flex: 1; overflow: hidden; }
.sf-header .nav span { padding: 10px 8px; opacity: 0.85; cursor: pointer; white-space: nowrap; }
.sf-header .nav span.active { opacity: 1; border-bottom: 2px solid #fff; font-weight: 600; }
.sf-tabs {
  background: #f3f3f3; padding: 0 16px; display: flex; gap: 20px;
  border-bottom: 1px solid #d8dde6; font-size: 12px;
}
.sf-tabs div { padding: 10px 4px; color: #54698d; border-bottom: 2px solid transparent; }
.sf-tabs div.active { color: #1B3A5C; border-bottom-color: #0070d2; font-weight: 600; }
.sf-body { padding: 18px; background: #f3f3f3; }
.sf-record-header {
  background: #fff; padding: 14px 18px; border-radius: 4px; border: 1px solid #dddbda;
  margin-bottom: 14px; display: flex; justify-content: space-between; align-items: center;
}
.sf-record-header .title { font-size: 11px; color: #706e6b; text-transform: uppercase; letter-spacing: 0.5px; }
.sf-record-header .name { font-size: 18px; color: #080707; font-weight: 600; margin-top: 2px; }
.sf-record-header .actions { display: flex; gap: 8px; }
.sf-btn {
  padding: 6px 14px; font-size: 12px; border: 1px solid #dddbda; background: #fff;
  border-radius: 4px; color: #0070d2; cursor: pointer;
}
.sf-btn.primary { background: #0070d2; color: #fff; border-color: #0070d2; }
.sf-btn.pulse { background: #C5A55A; color: #fff; border-color: #C5A55A; box-shadow: 0 0 0 3px rgba(197,165,90,0.3); font-weight: 600; }
.sf-highlights {
  background: #fff; padding: 16px; border-radius: 4px; border: 1px solid #dddbda;
  display: grid; grid-template-columns: repeat(4, 1fr); gap: 20px; margin-bottom: 14px;
}
.sf-highlights .field .label { font-size: 10px; color: #706e6b; text-transform: uppercase; letter-spacing: 0.3px; margin-bottom: 4px; }
.sf-highlights .field .value { font-size: 13px; color: #080707; font-weight: 500; }
.sf-path {
  background: #fff; border: 1px solid #dddbda; border-radius: 4px; padding: 10px 16px;
  display: flex; align-items: center; margin-bottom: 14px; font-size: 12px;
}
.sf-path .stage {
  flex: 1; padding: 6px 10px; background: #ecebea; margin-right: 2px; text-align: center;
  color: #706e6b; clip-path: polygon(0 0, 92% 0, 100% 50%, 92% 100%, 0 100%, 8% 50%);
}
.sf-path .stage.done { background: #04844b; color: #fff; }
.sf-path .stage.current { background: #0070d2; color: #fff; font-weight: 600; }

/* Convert dialog — INLINE (position:relative, not fixed) */
.sf-dialog {
  background: #fff; border-radius: 6px; border: 1px solid #dddbda; padding: 0;
  max-width: 100%; margin: 0 auto; box-shadow: 0 2px 8px rgba(0,0,0,0.15);
  position: relative;
}
.sf-dialog-header {
  background: #16325c; color: #fff; padding: 14px 20px; border-radius: 6px 6px 0 0;
  font-size: 15px; font-weight: 600; display: flex; justify-content: space-between;
}
.sf-dialog-body { padding: 20px; }
.sf-section {
  background: #fafaf9; border: 1px solid #dddbda; border-radius: 4px; padding: 14px; margin-bottom: 12px;
}
.sf-section h4 { margin: 0 0 10px 0; font-size: 13px; color: #1B3A5C; text-transform: uppercase; letter-spacing: 0.5px; }
.sf-field-row { display: grid; grid-template-columns: 160px 1fr; gap: 10px; margin-bottom: 8px; font-size: 12px; align-items: center; }
.sf-field-row .lbl { color: #3e3e3c; font-weight: 500; }
.sf-field-row .input { background: #fff; border: 1px solid #dddbda; padding: 6px 10px; border-radius: 3px; color: #080707; }
.sf-field-row .radio { display: flex; gap: 14px; font-size: 12px; }
.sf-field-row .radio label { display: flex; align-items: center; gap: 5px; }
.sf-field-row .radio input { accent-color: #0070d2; }
.sf-dialog-footer {
  padding: 14px 20px; border-top: 1px solid #dddbda; background: #f3f3f3;
  display: flex; justify-content: flex-end; gap: 10px; border-radius: 0 0 6px 6px;
}

/* Tables */
table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 13px; }
th { background: #1B3A5C; color: #fff; text-align: left; padding: 10px 12px; font-weight: 600; }
td { padding: 9px 12px; border-bottom: 1px solid #e0e0e5; }
tr:nth-child(even) td { background: #fafafa; }

/* TOC */
.toc { background: #F5F5F7; padding: 24px 30px; border-radius: 6px; margin-bottom: 30px; }
.toc h2 { margin-top: 0; border: none; }
.toc ol { margin: 0; padding-left: 20px; }
.toc li { margin: 6px 0; font-size: 14px; }

/* Result tags */
.result-tags { display: flex; gap: 10px; flex-wrap: wrap; margin: 12px 0; }
.result-tag {
  background: #fff; border: 1px solid #1B3A5C; color: #1B3A5C;
  padding: 8px 14px; border-radius: 20px; font-size: 12px; font-weight: 600;
}
.result-tag.account { border-color: #04844b; color: #04844b; }
.result-tag.contact { border-color: #0070d2; color: #0070d2; }
.result-tag.opp { border-color: #C5A55A; color: #9a7b2f; }

/* Captions */
.caption { text-align: center; font-size: 12px; color: #706e6b; font-style: italic; margin-top: -10px; margin-bottom: 20px; }

@media print { body { padding: 20px; } .cover { page-break-after: always; } }
"""

# Apply tenant brand colors (no-op when colors match the defaults above)
_redacted_CSS = (
    _redacted_CSS
    .replace("#1B3A5C", _PRIMARY)
    .replace("#2d5280", _PRIMARY_DARK)
    .replace("#C5A55A", _ACCENT)
)

_NEUTRAL_CSS = (_redacted_CSS
    .replace(_PRIMARY, "#2F3A4A")
    .replace(_PRIMARY_DARK, "#5B6778")
    .replace(_ACCENT, "#8B97A7")
    .replace("#F5F5F7", "#F7F6F3")
)


def _theme_css(prefs):
    return _NEUTRAL_CSS if (prefs or {}).get("brand_mode") == "neutral" else _redacted_CSS


def _build_sf_header_html(prefs):
    prefs = prefs or {}
    author = prefs.get("author") or prefs.get("display_name") or "AI"
    initials = "".join(part[:1] for part in author.split()[:2]).upper() or "AI"
    if prefs.get("salesforce_branding") != "generic":
        return _SF_HEADER_HTML
    return (
        "<div class=\"sf-valley-bar\">"
        "<span class=\"sf-valley-logo\">SALESFORCE</span>"
        "<input class=\"sf-search\" placeholder=\"Search Salesforce...\">"
        "<div class=\"sf-right\">"
        "<div class=\"sf-utils\">&#9734; &#43; &#8962; &#63; &#9881;</div>"
        f"<div class=\"sf-avatar\">{initials}<div class=\"notif\">3</div></div>"
        "</div></div>"
        "<div class=\"sf-header\">"
        "<div class=\"waffle\"><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span></div>"
        "<span class=\"app-name\">Lightning</span>"
        "<div class=\"nav\"><span>Home</span><span>Accounts</span><span>Contacts</span><span>Opportunities</span><span>Leads</span><span>Cases</span><span>Reports</span><span>Dashboards</span><span>More &#9662;</span></div>"
        "</div>"
    )

# Hardcoded Valley + Salesforce nav — replaces whatever Qwen3 wrote
_SF_HEADER_HTML = (
    '<div class="sf-valley-bar">'
    f'<span class="sf-valley-logo">{_SF_BRAND_LABEL}</span>'
    '<input class="sf-search" placeholder="Search...">'
    '<div class="sf-right">'
    '<div class="sf-utils">&#9734; &#43; &#8962; &#63; &#9881;</div>'
    '<div class="sf-avatar">AD<div class="notif">15</div></div>'
    '</div>'
    '</div>'
    '<div class="sf-header">'
    '<div class="waffle"><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span><span></span></div>'
    '<span class="app-name">Sales Assist</span>'
    '<div class="nav">'
    '<span>Home</span>'
    '<span>Accounts</span>'
    '<span>Contacts</span>'
    '<span>Opportunities</span>'
    '<span>Leads</span>'
    '<span>Tasks</span>'
    '<span>Dashboards</span>'
    '<span>TM Forecast/Sales Summary</span>'
    '<span>Quantity Forecasts</span>'
    '<span>Reports</span>'
    '<span>More &#9662;</span>'
    '</div>'
    '</div>'
)


def _inject_nav_bars(html_content, prefs=None):
    """Replace every sf-header, sf-nav, sf-valley-bar with the correct hardcoded nav."""
    # Kill sf-nav blocks
    html_content = re.sub(
        r'<div class="sf-nav">.*?(?=<div class="sf-header">|<div class="sf-tabs)',
        '', html_content, flags=re.DOTALL
    )
    # Kill any existing valley bars
    html_content = re.sub(r'<div class="sf-valley-bar">.*?</div>', '', html_content, flags=re.DOTALL)
    # Replace sf-header content with our Valley bar + SF nav
    html_content = re.sub(
        r'<div class="sf-header">.*?(?=<div class="sf-tabs|<div class="sf-body)',
        _build_sf_header_html(prefs) + '\n  ',
        html_content, flags=re.DOTALL
    )
    return html_content




def _restrict_sf_paths(content: str) -> str:
    """Keep Salesforce path visuals only in Lead/Opportunity contexts."""
    pattern = re.compile(r'<div class="sf-path">.*?</div>', re.DOTALL)

    def _replace(match):
        window = content[max(0, match.start() - 1400):match.end()].lower()
        allow_tokens = (
            '/lead/', '>lead<', ' lead ', 'lead record',
            '/opportunity/', '>opportunity<', ' opportunity ', 'opportunity record', ' opp '
        )
        return match.group(0) if any(token in window for token in allow_tokens) else ''

    return pattern.sub(_replace, content)


def _wrap_in_template(content, prefs=None):
    """Wrap content in the selected document template. Strips model-authored wrappers first."""
    prefs = prefs or get_doc_preferences("Operator")
    # Strip any CSS Qwen3 wrote — we provide all styling
    content = re.sub(r'<style[^>]*>.*?</style>', '', content, flags=re.DOTALL)

    # Extract body if Qwen3 wrote a full HTML doc
    body_match = re.search(r'<body[^>]*>(.*)</body>', content, re.DOTALL)
    if body_match:
        content = body_match.group(1).strip()
    else:
        # Strip any remaining HTML/head wrappers
        content = re.sub(r'<!DOCTYPE[^>]*>', '', content)
        content = re.sub(r'</?html[^>]*>', '', content)
        content = re.sub(r'<head[^>]*>.*?</head>', '', content, flags=re.DOTALL)
        content = re.sub(r'</?body[^>]*>', '', content)
        content = content.strip()

    # Remove stray sales-path UI from non-Lead/Opportunity mockups
    content = _restrict_sf_paths(content)

    # Extract title from first h1
    title_match = re.search(r'<h1[^>]*>(.*?)</h1>', content, re.DOTALL)
    title = re.sub(r'<[^>]+>', '', title_match.group(1)).strip() if title_match else 'Document'

    return f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8"><title>{title}</title>
<style>{_theme_css(prefs)}</style></head>
<body>
{content}
</body></html>"""

def _inject_toc_script(html_content):
    """Add IDs to headings for anchor links. Skip if Qwen3 already wrote a .toc div."""
    # If Qwen3 already wrote a TOC, just add heading IDs
    h2_count = len(re.findall(r'<h2', html_content))
    if h2_count < 3:
        return html_content

    # Add IDs to headings that don't have them
    counter = [0]
    def _add_id(match):
        counter[0] += 1
        tag = match.group(1)
        attrs = match.group(2) or ""
        text = match.group(3)
        if 'id=' not in attrs:
            attrs = f' id="sec-{counter[0]}"' + attrs
        return f'<{tag}{attrs}>{text}</{tag}>'

    html_content = re.sub(r'<(h[23])([^>]*)>(.*?)</\1>', _add_id, html_content, flags=re.DOTALL)
    return html_content


OUTPUT_DIR = os.path.expanduser(
    SETTINGS.get("outputs", {}).get("default_path", "~/Documents/Claude Outputs")
)


def _render_html_to_image(html_snippet, output_path, width=850):
    """Render an HTML snippet to a PNG image using headless Chrome."""
    import subprocess, tempfile
    CHROME = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
    # Write snippet to temp file with inline styles
    with tempfile.NamedTemporaryFile(suffix='.html', mode='w', delete=False, dir='/tmp') as f:
        f.write(html_snippet)
        tmp_path = f.name
    try:
        subprocess.run(
            [CHROME, "--headless", "--disable-gpu", "--hide-scrollbars", "--no-sandbox",
             f"--screenshot={output_path}", f"--window-size={width},2000",
             f"file://{tmp_path}"],
            capture_output=True, timeout=15,
        )
        return os.path.exists(output_path) and os.path.getsize(output_path) > 500
    except Exception as e:
        log.warning(f"Chrome screenshot failed: {e}")
        return False
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


def _extract_visual_sections(html_content):
    """Extract SVG flowcharts and SF mockups as renderable HTML snippets."""
    visuals = []
    # Extract style block for rendering context
    style = ""
    style_match = re.search(r'<style[^>]*>(.*?)</style>', html_content, re.DOTALL)
    if style_match:
        style = style_match.group(0)

    # Flowchart SVGs (with surrounding heading)
    svg_matches = re.finditer(r'(<h3[^>]*>Process Flow</h3>\s*<svg[^>]*>.*?</svg>)', html_content, re.DOTALL)
    for m in svg_matches:
        snippet = f'<!DOCTYPE html><html><head>{style}<style>body{{margin:20px;font-family:system-ui,sans-serif}}</style></head><body>{m.group(1)}</body></html>'
        visuals.append(("flowchart", snippet, m.start(), m.end()))

    # SF mockups (each mockup div with its heading)
    mockup_pattern = re.compile(r'(<h3[^>]*>\w+ Record Page</h3>\s*<div class="sf-mockup">.*?</div>\s*</div>\s*</div>)', re.DOTALL)
    for m in mockup_pattern.finditer(html_content):
        snippet = f'<!DOCTYPE html><html><head>{style}<style>body{{margin:20px;font-family:system-ui,sans-serif}}</style></head><body>{m.group(1)}</body></html>'
        visuals.append(("mockup", snippet, m.start(), m.end()))

    return visuals


def _html_to_docx(html_content, filepath, prefs=None):
    """Convert HTML to .docx using python-docx with proper formatting.
    Renders mockups and flowcharts as images via headless Chrome."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    prefs = prefs or get_doc_preferences("Operator")

    # Render visual sections (flowcharts + mockups) as images
    visuals = _extract_visual_sections(html_content)
    rendered_visuals = []
    for i, (vtype, snippet, start, end) in enumerate(visuals):
        img_path = f"/tmp/docx_visual_{i}.png"
        if _render_html_to_image(snippet, img_path, width=800 if vtype == "mockup" else 850):
            rendered_visuals.append((vtype, img_path))
            log.info(f"Rendered {vtype} to {img_path}")

    # Pre-process: strip scripts, SVG, mockups, TOC — text only for body
    clean = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
    clean = re.sub(r'<svg[^>]*>.*?</svg>', '', clean, flags=re.DOTALL)
    clean = re.sub(r'<div class="sf-mockup">.*?</div>\s*</div>\s*</div>', '', clean, flags=re.DOTALL)
    clean = re.sub(r'<div id="toc">.*?</div>', '', clean, flags=re.DOTALL)
    clean = re.sub(r'<style[^>]*>.*?</style>', '', clean, flags=re.DOTALL)
    # Strip the "Salesforce UI Reference" section headings (we'll rebuild with images)
    clean = re.sub(r'<h2[^>]*>Salesforce UI Reference</h2>.*', '', clean, flags=re.DOTALL)

    if prefs.get("brand_mode") == "neutral":
        primary_hex = "2F3A4A"
    else:
        primary_hex = _PRIMARY.lstrip("#")

    for style_name in ["Normal", "List Bullet"]:
        try:
            doc.styles[style_name].font.name = 'Segoe UI'
        except Exception:
            pass

    # Extract title
    title_match = re.search(r'<h1[^>]*>(.*?)</h1>', clean, re.DOTALL)
    if title_match:
        title = re.sub(r'<[^>]+>', '', title_match.group(1)).strip()
        if title:
            doc.add_heading(title, level=0)

    # Extract tables separately and convert them properly
    def _convert_table(table_html):
        """Convert an HTML table to a docx table."""
        rows_html = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL)
        if not rows_html:
            return
        # Parse all rows
        all_rows = []
        for row_html in rows_html:
            cells = re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', row_html, re.DOTALL)
            cells = [re.sub(r'<[^>]+>', '', c).strip() for c in cells]
            if cells:
                all_rows.append(cells)
        if not all_rows:
            return
        # Create table
        n_cols = max(len(r) for r in all_rows)
        table = doc.add_table(rows=len(all_rows), cols=n_cols)
        table.style = 'Table Grid'
        for i, row_data in enumerate(all_rows):
            for j, cell_text in enumerate(row_data):
                if j < n_cols:
                    table.rows[i].cells[j].text = cell_text
        # Bold header row
        if all_rows:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # Process content: split by tables first, then handle the rest
    # Replace tables with markers, process inline content, then insert tables
    table_pattern = re.compile(r'<table[^>]*>.*?</table>', re.DOTALL)
    tables = table_pattern.findall(clean)
    parts = table_pattern.split(clean)

    for idx, part in enumerate(parts):
        # Process non-table content
        sections = re.split(r'(<h[1-6][^>]*>.*?</h[1-6]>|<p[^>]*>.*?</p>|<li[^>]*>.*?</li>)', part, flags=re.DOTALL)

        for section in sections:
            section = section.strip()
            if not section:
                continue

            # Headings
            handled = False
            for level in range(1, 7):
                match = re.match(rf'<h{level}[^>]*>(.*?)</h{level}>', section, re.DOTALL)
                if match:
                    text = re.sub(r'<[^>]+>', '', match.group(1)).strip()
                    if text and text != 'Table of Contents':
                        doc.add_heading(text, level=min(level, 4))
                    handled = True
                    break

            if handled:
                continue

            # Paragraphs
            p_match = re.match(r'<p[^>]*>(.*?)</p>', section, re.DOTALL)
            if p_match:
                text = re.sub(r'<[^>]+>', '', p_match.group(1)).strip()
                if text:
                    doc.add_paragraph(text)
                continue

            # List items
            li_match = re.match(r'<li[^>]*>(.*?)</li>', section, re.DOTALL)
            if li_match:
                text = re.sub(r'<[^>]+>', '', li_match.group(1)).strip()
                if text:
                    doc.add_paragraph(text, style='List Bullet')
                continue

        # Insert table after this part (if there is one)
        if idx < len(tables):
            _convert_table(tables[idx])

    # Append rendered visuals (flowcharts + mockups) as images at end
    if rendered_visuals:
        doc.add_heading("Visual References", level=2)
        for vtype, img_path in rendered_visuals:
            try:
                doc.add_picture(img_path, width=Inches(6.0))
                doc.add_paragraph("")  # spacer
            except Exception as e:
                log.warning(f"Failed to insert {vtype} image: {e}")

    docx_path = re.sub(r'\.(html?|md|txt)$', '.docx', filepath)
    if not docx_path.endswith('.docx'):
        docx_path += '.docx'
    doc.save(docx_path)

    # Clean up temp images
    for _, img_path in rendered_visuals:
        try:
            os.unlink(img_path)
        except Exception:
            pass

    return docx_path


class DocumentOutputTool(BaseTool):
    name = "document_output"
    description = "Save documents, reports, guides, and artifacts to OneDrive Claude Outputs folder"

    def get_definitions(self):
        return [
            make_tool_def("save_document", "Save a document/report/guide to OneDrive Claude Outputs folder. Syncs automatically.",
                          {"filename": {"type": "string", "description": "File name with extension (e.g. 'Implementation Guide.html', 'report.html')"},
                           "content": {"type": "string", "description": "Full document content"},
                           "format": {"type": "string", "description": "Format: html, txt, csv, json (default: html)"},
                           "doc_type": {"type": "string", "description": "Either salesforce or corporate"},
                           "doc_profile": {"type": "string", "description": "Named document profile for template shaping"},
                           "user_id": {"type": "string", "description": "User ID for branding selection"}},
                          ["filename", "content"]),
            make_tool_def("list_outputs", "List files in the Claude Outputs folder.",
                          {}, []),
            make_tool_def("read_output", "Read a file from Claude Outputs folder.",
                          {"filename": {"type": "string"}}, ["filename"]),
        ]

    async def handle(self, tool_name, tool_input):
        if tool_name == "save_document":
            # Resilient key extraction — Mistral sometimes uses different key names
            filename = tool_input.get("filename") or tool_input.get("name") or tool_input.get("title") or tool_input.get("file_name") or "Document.html"
            content = tool_input.get("content") or tool_input.get("body") or tool_input.get("text") or tool_input.get("html") or ""
            fmt = tool_input.get("format", "html")
            user_id = (tool_input.get("user_id") or "Operator").lower()
            prefs = get_doc_preferences(user_id)
            doc_type = tool_input.get("doc_type", "").lower()
            # Auto-detect if not specified: check for SF mockup classes
            if not doc_type:
                doc_type = "salesforce" if (".sf-" in content or "sf-browser-bar" in content or "mockup" in content) else "corporate"
            if not content:
                return "Error: No content provided in tool call."

            os.makedirs(OUTPUT_DIR, exist_ok=True)

            filename = _sanitize_output_filename(filename, fmt)

            # Check for existing file with same base name (overwrite for iterative updates)
            now = datetime.now()
            timestamp = now.strftime("%Y-%m-%d %H%M")
            base_name = re.sub(r'^\d{4}-\d{2}-\d{2}\s+\d{4}\s*—\s*', '', filename)
            existing = [f for f in os.listdir(OUTPUT_DIR) if f.endswith(base_name)]
            if existing:
                filename = sorted(existing)[-1]
                log.info(f"Overwriting existing doc: {filename}")
            elif not filename.startswith("20"):
                filename = f"{timestamp} — {filename}"

            filepath = os.path.join(OUTPUT_DIR, filename)

            try:
                # Post-process HTML: wrap in template, conditionally inject SF nav bars
                if filepath.endswith('.html') or filepath.endswith('.htm'):
                    # Strip any <img> tags the model may have injected despite Rule 6
                    content = re.sub(r'<img\b[^>]*>', '', content, flags=re.IGNORECASE)
                    content = _wrap_in_template(content, prefs)
                    if doc_type == "salesforce":
                        content = _inject_nav_bars(content, prefs)
                    content = _inject_toc_script(content)

                # Write file — if locked by OneDrive, fall back to new timestamped name
                try:
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(content)
                except OSError as oe:
                    log.warning(f"File locked ({oe}), writing new version")
                    filename = f"{timestamp} — {base_name}"
                    filepath = os.path.join(OUTPUT_DIR, filename)
                    with open(filepath, "w", encoding="utf-8") as f:
                        f.write(content)
                log.info(f"Saved document ({fmt}): {filepath}")

                # Also save to local cache (no OneDrive locks)
                cache_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "doc_cache")
                os.makedirs(cache_dir, exist_ok=True)
                cache_path = os.path.join(cache_dir, filename)
                try:
                    with open(cache_path, "w", encoding="utf-8") as cf:
                        cf.write(content)
                except Exception as ce:
                    log.warning(f"Cache write failed: {ce}")

                result = f"Document saved to OneDrive: {filename}\nFormat: {fmt.upper()}\nPath: Claude Outputs/{filename}"

                # Auto-convert HTML to docx
                if filepath.endswith('.html') or filepath.endswith('.htm'):
                    try:
                        docx_path = _html_to_docx(content, filepath, prefs)
                        docx_name = os.path.basename(docx_path)
                        log.info(f"Auto-converted to docx: {docx_path}")
                        result += f"\nAlso saved as Word: {docx_name}"
                    except Exception as e:
                        log.warning(f"Auto-convert to docx failed: {e}")

                return result
            except Exception as e:
                return f"Error saving document: {e}"

        elif tool_name == "list_outputs":
            try:
                files = os.listdir(OUTPUT_DIR)
                files = [f for f in files if not f.startswith(".")]
                files.sort(reverse=True)
                if not files:
                    return "No files in Claude Outputs folder."
                return "\n".join(files[:30])
            except Exception as e:
                return f"Error listing outputs: {e}"

        elif tool_name == "read_output":
            filename = tool_input["filename"]
            filepath = os.path.join(OUTPUT_DIR, filename)
            try:
                with open(filepath, "r", encoding="utf-8") as f:
                    content = f.read()
                return content[:5000]
            except FileNotFoundError:
                return f"File not found: {filename}"
            except Exception as e:
                return f"Error reading file: {e}"

        else:
            return f"Unknown document tool: {tool_name}"
